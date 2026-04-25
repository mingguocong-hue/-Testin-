"""
utils.py — 招聘系统核心工具库
核心修复：
  1. extract_resume_from_bytes() —— 直接从 Streamlit 上传的字节流提取文本，
     彻底绕过文件路径/编码等磁盘读取问题（这是之前 resume_text 全为空的根因）
  2. _validate_against_source() —— 改为模糊子串匹配，不再误删合法技能
  3. jd_match_score() —— 优先把完整简历原文交给 AI 评分，不再仅靠表单关键词
  4. parse_resume_text() —— 提示词升级，覆盖更多中文简历格式
  5. save_to_excel() —— 新增获奖/GPA/个人特质/评分摘要等列，真正填满表格
"""

from __future__ import annotations

import re
import json
import os
from io import BytesIO
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv(override=False)
except ImportError:
    pass

import requests

try:
    from thefuzz import fuzz as _fuzz
    _FUZZ_OK = True
except ImportError:
    _FUZZ_OK = False

try:
    from docx import Document as _DocxDoc
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import pdfplumber as _pdfplumber
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ══════════════════════════════════════════════════════════════════════════════
# 一、基础验证
# ══════════════════════════════════════════════════════════════════════════════

def validate_phone(phone: str) -> bool:
    return bool(re.fullmatch(r"1[3-9]\d{9}", str(phone).strip()))


def validate_email(email: str) -> bool:
    return bool(re.fullmatch(
        r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}",
        str(email).strip()
    ))


# ══════════════════════════════════════════════════════════════════════════════
# 二、简历文本提取
# 核心修复：新增 extract_resume_from_bytes()，直接从内存字节流解析，
# 不依赖磁盘路径，彻底解决之前 resume_text 全为空的问题
# ══════════════════════════════════════════════════════════════════════════════

def _extract_docx_xml(file_bytes: bytes) -> str:
    """
    把 DOCX（本质是 ZIP）里所有 XML 文件中的 <w:t> 节点文字拼起来。
    覆盖范围：正文段落、表格单元格、文本框、形状、页眉、页脚。
    普通 doc.paragraphs 只能拿到正文段落，文本框模板会全部漏掉，这个方法解决该问题。
    """
    import zipfile
    import xml.etree.ElementTree as ET

    W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    W_R = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
    W_BR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"

    paragraphs: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(file_bytes)) as z:
            # 按文件名排序：document.xml 优先，再是 header/footer/textbox 等
            xml_files = sorted(
                [n for n in z.namelist() if n.startswith("word/") and n.endswith(".xml")],
                key=lambda x: (0 if "document" in x else 1, x)
            )
            for xml_name in xml_files:
                try:
                    root = ET.fromstring(z.read(xml_name))
                except ET.ParseError:
                    continue

                # 按段落分组：每个 <w:p> 作为一行
                for p_node in root.iter(W_P):
                    parts: list[str] = []
                    for t_node in p_node.iter(W_T):
                        txt = t_node.text or ""
                        if txt:
                            parts.append(txt)
                    line = "".join(parts).strip()
                    if line:
                        paragraphs.append(line)

    except zipfile.BadZipFile:
        return ""
    except Exception:
        return ""

    # 去除相邻重复行（DOCX 兼容性 XML 节点会导致同一行出现两次）
    deduped: list[str] = []
    for line in paragraphs:
        if not deduped or line != deduped[-1]:
            deduped.append(line)
    return "\n".join(deduped)


def extract_resume_from_bytes(file_bytes: bytes, suffix: str) -> tuple[str, str]:
    """
    从上传文件的字节流中提取纯文本。
    返回 (text, error_msg)，成功时 error_msg 为空字符串。

    在 app.py 里推荐这样调用（文件上传后立即提取，不依赖保存路径）：
        text, err = extract_resume_from_bytes(
            resume_file.getvalue(), Path(resume_file.name).suffix
        )
    """
    suffix = suffix.lower().strip()
    if not suffix.startswith("."):
        suffix = "." + suffix

    # ── TXT ──────────────────────────────────────────────────────────────────
    if suffix == ".txt":
        for enc in ("utf-8", "gbk", "utf-16", "gb2312"):
            try:
                return file_bytes.decode(enc), ""
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode(errors="replace"), ""

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if suffix == ".docx":
        # 方法一：直接解析 ZIP 内 XML，提取所有 <w:t> 节点
        # 覆盖范围：正文段落、表格、文本框、形状、页眉页脚（普通 doc.paragraphs 拿不到文本框）
        text = _extract_docx_xml(file_bytes)
        if text.strip():
            return text, ""

        # 方法二：python-docx 兜底（处理极少数非标准 XML 的情况）
        if _DOCX_OK:
            try:
                doc = _DocxDoc(BytesIO(file_bytes))
                lines: list[str] = []
                for p in doc.paragraphs:
                    t = p.text.strip()
                    if t:
                        lines.append(t)
                seen_cells: set[str] = set()
                for table in doc.tables:
                    for row in table.rows:
                        row_parts = []
                        for cell in row.cells:
                            ct = cell.text.strip()
                            if ct and ct not in seen_cells:
                                seen_cells.add(ct)
                                row_parts.append(ct)
                        if row_parts:
                            lines.append("  ".join(row_parts))
                text = "\n".join(lines)
                if text.strip():
                    return text, ""
            except Exception:
                pass

        return "", "DOCX 内容为空，请确认文件不是扫描图片；若使用特殊模板请另存为普通 DOCX 再上传"

    # ── PDF ──────────────────────────────────────────────────────────────────
    if suffix == ".pdf":
        if not _PDF_OK:
            return "", "pdfplumber 未安装，请执行: pip install pdfplumber"
        try:
            texts: list[str] = []
            with _pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if t and t.strip():
                        texts.append(t.strip())
            text = "\n".join(texts)
            if not text.strip():
                return "", "PDF 内容为空（扫描图片型简历无法提取文字，建议改用 DOCX 或 TXT）"
            return text, ""
        except Exception as e:
            return "", f"PDF 解析异常：{e}"

    return "", f"不支持的文件格式：{suffix}，请上传 TXT / DOCX / PDF"


def extract_resume_text(file_path) -> str:
    """兼容旧调用：从磁盘路径读取后调用 extract_resume_from_bytes。"""
    path = Path(file_path)
    if not path.exists():
        return ""
    try:
        text, _ = extract_resume_from_bytes(path.read_bytes(), path.suffix)
        return text
    except Exception:
        return ""


# ══════════════════════════════════════════════════════════════════════════════
# 三、Claude API
# ══════════════════════════════════════════════════════════════════════════════

def _call_claude(system_prompt: str, user_prompt: str, max_tokens: int = 1500) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY 未配置")
    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        },
        json={
            "model": "claude-sonnet-4-20250514",
            "max_tokens": max_tokens,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        },
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()["content"][0]["text"]


def _strip_json(raw: str) -> str:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?", "", raw).strip()
    raw = re.sub(r"```$", "", raw).strip()
    return raw


# ══════════════════════════════════════════════════════════════════════════════
# 四、简历解析（AI 优先 → 正则兜底）
# ══════════════════════════════════════════════════════════════════════════════

def parse_resume_text(text: str) -> dict:
    """从简历原文提取结构化字段。AI 可用时走 Claude，否则走正则兜底。"""
    if not text or len(text.strip()) < 10:
        return {}

    try:
        system = (
            "你是专业的中文简历信息提取助手。"
            "严格只返回纯 JSON，不含任何 Markdown 标记、注释或额外文字。"
            "字段不存在时设为 null 或空列表，不要捏造不在原文中的信息。"
        )
        user = (
            "请从以下简历文本中提取所有可识别的信息，按如下 JSON 结构返回：\n\n"
            "{\n"
            '  "name": "真实姓名2~4个字，原文没有则null",\n'
            '  "phone": "11位手机号，原文没有则null",\n'
            '  "email": "邮箱地址，原文没有则null",\n'
            '  "school": "最高学历学校全称，原文没有则null",\n'
            '  "major": "专业名称，原文没有则null",\n'
            '  "degree": "学历层次如本科/硕士/博士/大专，原文没有则null",\n'
            '  "english_level": "英语水平如CET-6 560分/雅思7.0/四级通过，原文没有则null",\n'
            '  "skills_programming": ["只写原文出现的编程语言、框架、数据库、工具"],\n'
            '  "skills_office": ["只写原文出现的办公软件如Word/Excel/PPT/PS/PR等"],\n'
            '  "skills_ai_tools": ["只写原文出现的AI工具如ChatGPT/Claude/Midjourney等"],\n'
            '  "internship": "实习经历摘要含公司+岗位+时间，没有则null",\n'
            '  "campus_leadership": "校园干部/学生会/社团职务，没有则null",\n'
            '  "personal_traits": ["从自我评价/个人简介中提取的性格特质词"],\n'
            '  "awards": "获奖经历摘要，没有则null",\n'
            '  "gpa": "GPA数值或成绩排名，没有则null"\n'
            "}\n\n"
            f"简历原文：\n{text[:4000]}"
        )
        raw = _call_claude(system, user, max_tokens=1500)
        parsed = json.loads(_strip_json(raw))
        return _validate_against_source(parsed, text)
    except Exception:
        pass

    return _regex_fallback(text)


def _validate_against_source(parsed: dict, source: str) -> dict:
    """
    防幻觉校验：
    - 手机/邮箱：必须出现在原文，否则清空并发出警告
    - 技能列表：宽松子串匹配（兼容大小写/空格差异），不再误删合法技能
    """
    result = dict(parsed)
    src_lo = source.lower()

    # 手机号严格校验
    phone = result.get("phone")
    if phone:
        ps = re.sub(r"\D", "", str(phone))
        if not re.fullmatch(r"1[3-9]\d{9}", ps) or ps not in source:
            result["phone"] = None
            result["_warn_phone"] = "AI 提取的手机号原文中未找到，已清空"

    # 邮箱严格校验
    email = result.get("email")
    if email and (not validate_email(str(email)) or str(email).lower() not in src_lo):
        result["email"] = None
        result["_warn_email"] = "AI 提取的邮箱原文中未找到，已清空"

    # 技能列表：宽松校验（核心词出现在原文即保留）
    for key in ("skills_programming", "skills_office", "skills_ai_tools", "personal_traits"):
        items = result.get(key) or []
        if not isinstance(items, list):
            result[key] = []
            continue
        kept = []
        for skill in items:
            skill_str = str(skill).strip()
            if not skill_str:
                continue
            # 按常见分隔符拆分，任意有意义的部分（>1字符）出现在原文即保留
            parts = re.split(r"[\s/\-_.、，,·]", skill_str)
            meaningful = [p for p in parts if len(p) > 1]
            if not meaningful:
                if skill_str.lower() in src_lo:
                    kept.append(skill_str)
            elif any(p.lower() in src_lo for p in meaningful):
                kept.append(skill_str)
        result[key] = kept

    return result


# ── 正则兜底（无 API 时使用）────────────────────────────────────────────────

def _regex_fallback(text: str) -> dict:
    d: dict = {}
    d["name"] = (_re1(r"姓\s*名\s*[：:]\s*(\S{2,4})", text)
                 or _re1(r"^([^\s\n]{2,4})\s*(?:男|女|\d{4})", text, re.M))
    d["phone"] = _re1(r"(1[3-9]\d{9})", text)
    d["email"] = _re1(r"([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})", text)
    d["school"] = (_re1(r"(?:毕业院校|就读院校|学\s*校|院校)\s*[：:]\s*([^\n，,]{2,20})", text)
                   or _re1(r"([\u4e00-\u9fa5]{2,8}(?:大学|学院|理工|财经|师范|农业|医科|外国语)\S*)", text))
    d["major"] = _re1(r"专\s*业\s*[：:]\s*([^\n，,]{2,20})", text)
    d["degree"] = (_re1(r"(?:学\s*历|学位)\s*[：:]\s*([^\n，,]{2,10})", text)
                   or _re1(r"(本科|硕士|博士|研究生|大专|在读)", text))
    d["english_level"] = (
        _re1(r"(CET[-\s]?6[\s,，]*\d{3,}\s*分?)", text, re.I) or
        _re1(r"(英语六级[\s,，]*\d{3,}\s*分?)", text) or
        _re1(r"(六级\s*(?:通过|\d{3,}\s*分)?)", text) or
        _re1(r"(CET[-\s]?4[\s,，]*\d{3,}\s*分?)", text, re.I) or
        _re1(r"(四级\s*(?:通过|\d{3,}\s*分)?)", text) or
        _re1(r"((?:雅思|IELTS|托福|TOEFL)[^\n，,]{0,20})", text, re.I) or
        _re1(r"(能(?:流利|熟练)(?:使用|阅读|交流)[^\n，,]{0,15}英语[^\n，,]{0,15})", text)
    )
    d["skills_programming"] = _find_kws(text, [
        "Python", "Java", "C\\+\\+", "C#", "Go", "Golang", "JavaScript", "TypeScript",
        "PHP", "Swift", "Kotlin", "MySQL", "PostgreSQL", "MongoDB", "Redis", "SQLite",
        "Spring Boot", "Django", "Flask", "FastAPI", "React", "Vue",
        "Matlab", "R语言", "爬虫", "机器学习", "深度学习",
        "TensorFlow", "PyTorch", "Pandas", "NumPy", "Scikit",
        "Docker", "Kubernetes", "Linux", "Git", "Shell", "Hadoop", "Spark",
    ])
    d["skills_office"] = _find_kws(text, [
        "Word", "Excel", "PPT", "PowerPoint", "WPS", "Visio", "PS", "PR", "Photoshop",
    ])
    d["skills_ai_tools"] = _find_kws(text, [
        "ChatGPT", "Claude", "DeepSeek", "豆包", "文心一言", "讯飞星火",
        "Kimi", "通义千问", "Gemini", "Copilot", "Midjourney",
        "Stable Diffusion", "Cursor",
    ])

    intern_blk = _extract_block(
        text, [r"实习经[历验]", r"工作经[历验]", r"工作经验"],
        [r"校园经[历验]", r"项目经[历验]", r"教育背景", r"技能", r"获奖"])
    d["internship"] = intern_blk[:200].strip() if intern_blk else None

    lead_blk = _extract_block(
        text, [r"校园经[历验]", r"学生工作", r"社团经[历验]", r"在校经[历验]"],
        [r"实习经[历验]", r"工作经[历验]", r"技能", r"获奖", r"自我评价"])
    lead_kws = ["班长", "团支书", "学生会", "社长", "部长", "队长", "干部", "负责人", "主席", "会长"]
    if lead_blk:
        d["campus_leadership"] = lead_blk[:150].strip()
    else:
        found = [k for k in lead_kws if k in text]
        d["campus_leadership"] = "、".join(found) if found else None

    trait_blk = _extract_block(
        text, [r"自我评价", r"个人评价", r"个人简介", r"个人优势", r"自我介绍"],
        [r"实习", r"校园", r"技能", r"获奖"])
    trait_kws = ["积极", "主动", "乐观", "责任心", "团队", "协作", "沟通", "学习能力",
                 "抗压", "创新", "细心", "严谨", "高效", "适应能力", "执行力", "领导力", "好奇心", "上进"]
    if trait_blk:
        found = [k for k in trait_kws if k in trait_blk]
        d["personal_traits"] = found if found else [trait_blk[:60].strip()]
    else:
        d["personal_traits"] = [k for k in trait_kws if k in text]

    d["awards"] = _extract_block(
        text, [r"获奖", r"荣誉", r"奖项"],
        [r"实习", r"校园", r"技能", r"教育", r"自我"])[:100].strip() or None
    d["gpa"] = _re1(r"GPA\s*[：:]\s*([\d.]+)", text, re.I)
    return d


def _re1(pattern: str, text: str, flags=0):
    m = re.search(pattern, text, flags)
    return m.group(1) if m else None


def _find_kws(text: str, keywords: list) -> list:
    found, seen = [], set()
    for kw in keywords:
        pattern = kw.replace("+", r"\+")
        if re.search(pattern, text, re.I) and kw.lower() not in seen:
            found.append(kw)
            seen.add(kw.lower())
    return found


def _extract_block(text: str, starts: list, ends: list) -> str:
    m = re.search("|".join(starts), text)
    if not m:
        return ""
    pos = m.end()
    m2 = re.search("|".join(ends), text[pos:])
    return text[pos: pos + m2.start()] if m2 else text[pos: pos + 500]


# ══════════════════════════════════════════════════════════════════════════════
# 五、按岗位定制评分维度 + JD 综合评分
# ══════════════════════════════════════════════════════════════════════════════

# ── 每个岗位的评分维度配置（维度key → label + 满分）────────────────────────────
# 所有岗位的各维度分数之和必须 = 100
POSITION_PROFILES: dict[str, dict] = {

    "管培生": {
        "personal_traits":       {"label": "综合素质与领导潜力",                 "max": 30},
        "comprehensive_ability": {"label": "综合能力（沟通/解决问题/学习/抗压等）", "max": 15},
        "awards":                {"label": "荣誉与奖励",                          "max":  5},
        "office":                {"label": "办公软件",                             "max": 10},
        "english":               {"label": "英语能力",                             "max": 10},
        "ai_tools":              {"label": "AI工具使用",                           "max": 10},
        "internship":            {"label": "行业实习与实践经历",                   "max": 10},
        "leadership":            {"label": "校园领导经历",                         "max": 10},
    },

    "AI测试员": {
        "personal_traits": {"label": "学习能力与细致程度",   "max": 20},
        "programming":      {"label": "编程与测试技能",       "max": 25},
        "ai_tools":         {"label": "AI工具使用与理解",     "max": 25},
        "english":          {"label": "英语能力",             "max": 10},
        "internship":       {"label": "相关实习经历",         "max": 10},
        "office":           {"label": "办公软件",             "max":  5},
        "leadership":       {"label": "校园领导经历",         "max":  5},
    },

    "后台技术管理员": {
        "personal_traits": {"label": "个人能力与管理潜力",   "max": 20},
        "programming":      {"label": "编程技能",             "max": 30},
        "ai_tools":         {"label": "AI工具使用",           "max": 15},
        "english":          {"label": "英语能力",             "max": 10},
        "internship":       {"label": "相关实习经历",         "max": 15},
        "office":           {"label": "办公软件",             "max":  5},
        "leadership":       {"label": "校园领导经历",         "max":  5},
    },

    "运营岗位": {
        "personal_traits": {"label": "沟通表达与创意能力",   "max": 30},
        "office":           {"label": "办公软件",             "max": 10},
        "english":          {"label": "英语能力",             "max": 10},
        "ai_tools":         {"label": "AI工具使用",           "max": 10},
        "internship":       {"label": "实习与运营相关经历",   "max": 10},
        "teamwork":         {"label": "团队与活动经历",       "max": 20},
        "leadership":       {"label": "校园领导经历",         "max": 10},
    },

    # 兜底：岗位名不在上面列表时使用
    "_default": {
        "personal_traits": {"label": "个人性格与能力",       "max": 30},
        "programming":      {"label": "编程技能",             "max": 20},
        "office":           {"label": "办公软件",             "max": 10},
        "english":          {"label": "英语能力",             "max": 10},
        "ai_tools":         {"label": "AI工具使用",           "max": 10},
        "internship":       {"label": "相关实习经历",         "max": 10},
        "leadership":       {"label": "校园领导经历",         "max": 10},
    },
}

# 兼容旧 admin.py 直接引用 SCORE_DIMENSIONS 的地方
SCORE_DIMENSIONS = POSITION_PROFILES["_default"]


def get_score_profile(position: str) -> dict:
    """根据岗位名返回对应的评分维度配置，未匹配时返回默认配置。"""
    return POSITION_PROFILES.get(position, POSITION_PROFILES["_default"])


# ── 各维度的 AI 评分说明（动态拼进 prompt）───────────────────────────────────
_DIM_CRITERIA: dict[str, str] = {
    "personal_traits": (
        "从自我评价/项目描述/校园经历中识别积极主动、团队协作、沟通能力、领导力、"
        "抗压能力、学习能力、严谨细致等特质，每个明确特质+4~6分；"
        "有1~2个团队合作项目或经历即可得10分以上；"
        "有2~3项获奖/荣誉额外加5分，满分封顶"
    ),
    "comprehensive_ability": (
        "综合评估以下能力，每项有明确体现+2分：\n"
        "  团队协作（参与或负责团队项目）、领导能力（带领他人/项目负责人）、\n"
        "  沟通能力（有跨部门/对外沟通经历）、解决问题能力（描述了解决实际问题的案例）、\n"
        "  学习能力（快速掌握新知识/技能的表现）、抗压能力（高强度任务/比赛经历）、\n"
        "  积极主动（自发推进项目/主动承担责任）；\n"
        "  有1~2项明确体现即可达到10分以上，不要求面面俱到"
    ),
    "teamwork": (
        "识别社团/志愿/公益/支教/竞赛/团队项目等集体活动，"
        "活动丰富且有实质贡献者给高分，仅列名不说内容给低分"
    ),
    "awards": (
        "识别奖学金/竞赛获奖/荣誉称号，国家级>省级>校级；"
        "有1项奖项→60%满分，有2~3项→满分，无奖项→0分"
    ),
    "programming": (
        "识别编程语言/框架/数据库/工具，技能越多样且有实际项目经验分越高，"
        "有 GitHub/作品集额外加分"
    ),
    "office": (
        "Word/PPT/Excel/WPS/PS/PR 等，熟练使用2个以上得满分"
    ),
    "english": (
        "六级425+→满分，六级通过→70%，四级→40%，雅思/托福→80~100%，"
        "简历有英文段落表明有基础→30~50%，无任何英语信息→0分"
    ),
    "ai_tools": (
        "ChatGPT/Claude/DeepSeek/Midjourney/Copilot等，"
        "每个有实际使用证据的工具+20~30%满分，仅列名不说用法打折"
    ),
    "internship": (
        "分两部分各占50%满分：\n"
        "  ① 有任何实习或实质性项目经历→50%满分；无→0分\n"
        "  ② 行业相关性（根据岗位判断，见下方岗位说明）→另外50%满分"
    ),
    "leadership": (
        "无→0，一般职务（委员/干事）→50%，"
        "班长/学生会主席/社团负责人→70~100%，任期越长越高"
    ),
}


def jd_match_score(candidate: dict, resume_text: str = "") -> dict:
    """
    按候选人意向岗位选择对应评分维度，总分100分。
    AI 评分失败时自动降级到规则评分。
    """
    position = candidate.get("target_position", "") or ""
    profile  = get_score_profile(position)

    summary_fields = {
        "name", "gender", "school", "major", "degree", "target_position",
        "english_level", "skills_programming", "skills_office", "skills_ai_tools",
        "internship", "campus_leadership", "personal_traits", "awards", "gpa", "strengths",
    }
    summary = {k: v for k, v in candidate.items() if k in summary_fields and v}

    # 动态构建评分维度说明
    dim_lines = []
    for i, (dim_key, meta) in enumerate(profile.items(), 1):
        criteria = _DIM_CRITERIA.get(dim_key, "根据简历内容综合判断")
        dim_lines.append(f"{i}. {meta['label']}（满分{meta['max']}分）：{criteria}")

    # 动态构建返回 JSON 格式模板
    json_template = "{" + ",".join(
        f'"{k}":{{"score":0,"reason":"具体理由"}}' for k in profile
    ) + ',"summary":"整体一句话评价"}'

    # 实习行业相关性标准：按岗位区分
    if position == "运营岗位":
        internship_relevance = (
            "【实习行业相关性标准（运营岗位）】\n"
            "  ② 与运营/宣传/市场营销/新媒体/活动策划/内容创作/品牌推广相关经历→另外50%满分；"
            "有实习但方向不相关→25%满分"
        )
    else:
        internship_relevance = (
            "【实习行业相关性标准】\n"
            "  ② 与AI/编程/技术/数据/产品/管理相关实习，或大厂（腾讯/阿里/字节/百度/华为等）实习→另外50%满分；"
            "有实习但方向不太相关→25%满分"
        )

    try:
        system = "你是专业招聘评分专家，只返回纯 JSON，不含任何 Markdown 或额外文字。"
        user = (
            f"岗位：{position or '未指定'}\n\n"
            "请仔细阅读候选人的简历原文，按以下维度打分（总分100分）。"
            "每个维度给整数分和1~2句具体理由，必须基于简历原文，不要凭空猜测。\n\n"
            "【评分维度】\n" + "\n".join(dim_lines) + "\n\n"
            + internship_relevance + "\n\n"
            f"候选人基本信息（供参考）：\n{json.dumps(summary, ensure_ascii=False)}\n\n"
            f"简历原文（主要评分依据）：\n{resume_text[:3500] if resume_text else '无，请根据基本信息尽力评分'}\n\n"
            f"返回格式（严格按此）：\n{json_template}"
        )
        raw    = _call_claude(system, user, max_tokens=1200)
        result = json.loads(_strip_json(raw))

        scores, total = {}, 0
        for dim_key, meta in profile.items():
            dim_d = result.get(dim_key, {})
            if not isinstance(dim_d, dict):
                dim_d = {}
            s = max(0, min(int(dim_d.get("score", 0)), meta["max"]))
            scores[dim_key] = {
                "score": s, "max": meta["max"],
                "label": meta["label"], "reason": dim_d.get("reason", ""),
            }
            total += s
        return {
            "scores": scores, "total": total,
            "summary": result.get("summary", ""),
            "position": position,
        }

    except Exception:
        pass

    return _rule_score(candidate, resume_text, profile, position)


def _rule_score(candidate: dict, resume_text: str,
                profile: dict | None = None, position: str = "") -> dict:
    """规则兜底评分，按 profile 维度打分。"""
    if profile is None:
        profile = POSITION_PROFILES["_default"]
    full = resume_text + " " + json.dumps(candidate, ensure_ascii=False)
    scores: dict = {}

    trait_kws = ["积极", "主动", "乐观", "责任心", "团队", "协作", "沟通", "学习能力",
                 "抗压", "创新", "细心", "严谨", "领导力", "好奇心", "上进", "执行力"]
    big_cos = ["腾讯", "阿里", "字节", "百度", "华为", "京东", "美团", "滴滴",
               "微软", "谷歌", "Meta", "Amazon", "网易", "小米", "蚂蚁", "bilibili"]

    for dim_key, meta in profile.items():
        mmax = meta["max"]
        label = meta["label"]

        if dim_key == "personal_traits":
            traits = list(candidate.get("personal_traits") or [])
            traits += [k for k in trait_kws if k in full and k not in traits]
            for s in (candidate.get("strengths") or []):
                for kw in trait_kws:
                    if kw in str(s) and kw not in traits:
                        traits.append(kw)
            # 每个特质 5 分 → 2个团队相关词（"团队"+"协作"）就能达到 10 分以上
            base = min(len(traits) * 5, mmax)
            # 奖项加成并入 personal_traits：2~3项奖项即得满奖项加分
            awards_text = candidate.get("awards") or ""
            award_kws_pt = ["奖学金", "一等奖", "二等奖", "三等奖", "优秀", "国家级", "省级", "冠军", "荣誉"]
            award_found_pt = [k for k in award_kws_pt if k in (awards_text + full)]
            awards_bonus = min(len(award_found_pt) * 3, 8) if award_found_pt else (3 if awards_text else 0)
            s = min(base + awards_bonus, mmax)
            r_parts = [f"识别特质：{', '.join(traits[:6]) or '暂无'}"]
            if award_found_pt:
                r_parts.append(f"奖项加成+{awards_bonus}分")
            r = "；".join(r_parts)

        elif dim_key == "comprehensive_ability":
            # 综合能力：每项能力关键词命中 +2分，1~2项体现即可达10分以上
            comp_kws = {
                "团队协作": ["团队", "协作", "合作", "一起", "共同"],
                "领导能力": ["领导", "负责人", "带领", "主导", "组织"],
                "沟通能力": ["沟通", "协调", "对接", "交流", "汇报"],
                "解决问题": ["解决", "优化", "改进", "攻克", "突破", "方案"],
                "学习能力": ["学习", "掌握", "自学", "快速", "研究"],
                "抗压能力": ["抗压", "压力", "紧张", "高强度", "竞赛", "比赛"],
                "积极主动": ["积极", "主动", "自发", "推进", "上进"],
            }
            hit_dims = []
            for dim_name, kws in comp_kws.items():
                if any(k in full for k in kws):
                    hit_dims.append(dim_name)
            s = min(len(hit_dims) * 2 + 2, mmax)   # 基础2分，每项+2，上限15
            r = f"识别能力维度：{', '.join(hit_dims) or '暂无'}"

        elif dim_key == "teamwork":
            team_kws = ["社团", "志愿", "公益", "支教", "竞赛", "团队", "活动", "组织", "协会"]
            found = [k for k in team_kws if k in full]
            s = min(len(found) * (mmax // 6 + 1), mmax)
            r = f"识别活动：{', '.join(found) or '暂未识别'}"

        elif dim_key == "awards":
            # 独立 awards 维度（管培生等有此维度）：2~3项即满分
            award_text = candidate.get("awards") or ""
            award_kws = ["奖学金", "一等奖", "二等奖", "三等奖", "优秀", "国家级", "省级", "冠军", "荣誉"]
            found = [k for k in award_kws if k in (award_text + full)]
            s = min(len(found) * 3, mmax) if found else (int(mmax * 0.3) if award_text else 0)
            r = f"奖项：{award_text[:60] or '暂未识别'}"

        elif dim_key == "programming":
            prog = candidate.get("skills_programming") or []
            s = min(len(prog) * (mmax // 5 + 1), mmax)
            r = f"技能：{', '.join(prog) or '暂未识别'}"

        elif dim_key == "office":
            office = candidate.get("skills_office") or []
            s = min(len(office) * (mmax // 3 + 1), mmax)
            r = f"软件：{', '.join(office) or '暂未识别'}"

        elif dim_key == "english":
            eng = candidate.get("english_level") or ""
            if re.search(r"六级|CET-?6", eng, re.I):
                nums = re.findall(r"\d{3,}", eng)
                s = mmax if (nums and int(nums[0]) >= 425) else int(mmax * 0.7)
                r = f"英语六级：{eng}"
            elif re.search(r"四级|CET-?4", eng, re.I):
                s, r = int(mmax * 0.4), "英语四级"
            elif re.search(r"雅思|IELTS|托福|TOEFL", eng, re.I):
                s, r = int(mmax * 0.9), f"国际英语考试：{eng}"
            elif eng:
                s, r = int(mmax * 0.3), eng
            else:
                s, r = 0, "未识别到英语水平"

        elif dim_key == "ai_tools":
            ai = candidate.get("skills_ai_tools") or []
            s = min(len(ai) * (mmax // 4 + 1), mmax)
            r = f"工具：{', '.join(ai) or '暂未识别'}"

        elif dim_key == "internship":
            intern_text = candidate.get("internship") or ""
            half = mmax // 2  # 5分（假设 mmax=10）

            # ① 有任何实习/项目经历 → 5分
            part1 = half if intern_text else 0

            # ② 行业相关性 → 5分，按岗位区分
            part2 = 0
            if intern_text:
                if position == "运营岗位":
                    ops_kws = ["运营", "宣传", "市场", "推广", "新媒体", "策划",
                               "活动", "内容", "品牌", "营销", "公关", "传播", "编辑"]
                    if any(k in intern_text for k in ops_kws):
                        part2 = half
                    elif any(b in intern_text for b in big_cos):
                        part2 = half // 2
                    else:
                        part2 = half // 4
                elif position == "管培生":
                    # 管培生：AI/编程/大厂/行业相关实习才加行业分
                    mgmt_kws = ["AI", "人工智能", "编程", "技术", "数据", "算法",
                                "产品", "管理", "咨询", "金融", "运营", "研究"]
                    if any(b in intern_text for b in big_cos):
                        part2 = half
                    elif any(k in intern_text for k in mgmt_kws):
                        part2 = half
                    else:
                        part2 = half // 4  # 有实习但方向不相关，少量分
                else:
                    tech_kws = ["AI", "编程", "技术", "开发", "数据", "算法",
                                "产品", "管理", "测试", "运维", "机器学习", "深度学习"]
                    if any(b in intern_text for b in big_cos):
                        part2 = half
                    elif any(k in intern_text for k in tech_kws):
                        part2 = half
                    else:
                        part2 = half // 4

            s = part1 + part2
            r_parts = []
            if part1: r_parts.append(f"有实习经历(+{part1}分)")
            if part2 == half: r_parts.append(f"行业相关(+{part2}分)")
            elif part2: r_parts.append(f"行业相关性一般(+{part2}分)")
            elif intern_text: r_parts.append("行业相关性待确认")
            r = "；".join(r_parts) if r_parts else "未识别到实习经历"

        elif dim_key == "leadership":
            lead = candidate.get("campus_leadership") or ""
            s, r = (int(mmax * 0.7), f"领导经历：{lead[:60]}") if lead else (0, "未识别到校园领导经历")

        else:
            s, r = 0, "规则评分未覆盖此维度"

        scores[dim_key] = {"score": s, "max": mmax, "label": label, "reason": r}

    total = sum(v["score"] for v in scores.values())
    return {
        "scores": scores, "total": total,
        "summary": f"规则评分（API未配置或调用失败），{position or '默认'}岗位，总分 {total}/100",
        "position": position,
    }



# ══════════════════════════════════════════════════════════════════════════════
# 六·五、图片简历识别（JPG / PNG → 调用 Claude Vision 提取文字）
# ══════════════════════════════════════════════════════════════════════════════

def extract_image_resume(file_bytes: bytes, suffix: str) -> tuple[str, str]:
    """
    用 Claude Vision API 识别图片型简历（JPG/PNG）中的文字。
    返回 (text, error_msg)，成功时 error_msg 为空。
    需要配置 ANTHROPIC_API_KEY，否则返回提示信息。
    """
    import base64

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return "", "图片简历识别需要配置 ANTHROPIC_API_KEY，当前未配置"

    suffix = suffix.lower().strip().lstrip(".")
    media_type = "image/jpeg" if suffix in ("jpg", "jpeg") else "image/png"
    b64_data = base64.standard_b64encode(file_bytes).decode("utf-8")

    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 2000,
                "messages": [{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": b64_data,
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                "这是一份简历图片，请提取图片中所有可见的文字内容，"
                                "保留原始排版结构和换行，分区域输出（如教育背景、工作经历、技能等）。"
                                "只输出文字内容本身，不要加任何解释或说明。"
                            ),
                        },
                    ],
                }],
            },
            timeout=60,
        )
        resp.raise_for_status()
        text = resp.json()["content"][0]["text"].strip()
        if not text:
            return "", "图片中未识别到文字（图片可能模糊或分辨率过低）"
        return text, ""
    except Exception as e:
        return "", f"图片识别失败：{e}"


# ══════════════════════════════════════════════════════════════════════════════
# 六·六、简历与表单信息一致性校验
# ══════════════════════════════════════════════════════════════════════════════

def check_resume_consistency(
    form_name: str,
    form_phone: str,
    form_email: str,
    parsed: dict,
) -> list[str]:
    """
    对比表单填写的姓名/手机/邮箱与简历中 AI/正则提取的结果。
    发现不一致时返回提示列表，让用户确认或修正。
    """
    issues: list[str] = []

    p_name  = (parsed.get("name")  or "").strip()
    p_phone = re.sub(r"\D", "", parsed.get("phone") or "")
    p_email = (parsed.get("email") or "").strip().lower()

    f_name  = form_name.strip()
    f_phone = re.sub(r"\D", "", form_phone.strip())
    f_email = form_email.strip().lower()

    # 姓名：简历识别到且与表单不一致
    if p_name and f_name and p_name != f_name:
        issues.append(
            f"**姓名不一致**：您在表单填写的是「{f_name}」，"
            f"但简历中识别到的是「{p_name}」，请确认哪个是正确姓名。"
        )

    # 手机号：只比较11位数字部分
    if p_phone and f_phone and p_phone != f_phone:
        issues.append(
            f"**手机号不一致**：表单填写「{form_phone.strip()}」，"
            f"简历中识别到「{parsed.get('phone')}」，请确认是否填写正确。"
        )

    # 邮箱：忽略大小写比较
    if p_email and f_email and p_email != f_email:
        issues.append(
            f"**邮箱不一致**：表单填写「{form_email.strip()}」，"
            f"简历中识别到「{parsed.get('email')}」，请确认是否填写正确。"
        )

    return issues

def duplicate_check(new_candidate: dict, all_candidates: list) -> list:
    rows = []
    new_phone = str(new_candidate.get("phone") or "")
    new_email = str(new_candidate.get("email") or "").lower()
    new_name = str(new_candidate.get("name") or "")

    for old in all_candidates:
        old_phone = str(old.get("phone") or "")
        old_email = str(old.get("email") or "").lower()
        old_name = str(old.get("name") or "")
        if new_phone and new_phone == old_phone and new_name == old_name:
            continue
        reasons, exact = [], 0
        if new_phone and new_phone == old_phone:
            exact = 100
            reasons.append("手机号一致")
        if new_email and new_email == old_email:
            exact = 100
            reasons.append("邮箱一致")
        if _FUZZ_OK:
            fuzzy = (0.6 * _fuzz.ratio(new_name, old_name) +
                     0.4 * _fuzz.ratio(str(new_candidate.get("school") or ""),
                                       str(old.get("school") or "")))
        else:
            fuzzy = 100.0 if new_name == old_name else 0.0
        score = max(exact, fuzzy)
        if score >= 75:
            rows.append({
                "候选人姓名": old_name, "电话": old_phone,
                "相似度": round(score, 1),
                "判断": "重复，建议合并" if score >= 95 else "疑似重复，请人工确认",
                "原因": "、".join(reasons) or "姓名/学校模糊匹配",
            })
    return sorted(rows, key=lambda x: -x["相似度"])


# ══════════════════════════════════════════════════════════════════════════════
# 七、候选人更新
# ══════════════════════════════════════════════════════════════════════════════

def update_candidate(new_candidate: dict, old_candidates: list):
    new_phone = str(new_candidate.get("phone") or "")
    new_email = str(new_candidate.get("email") or "").lower()
    for i, old in enumerate(old_candidates):
        if (new_phone and new_phone == str(old.get("phone") or "")) or \
           (new_email and new_email == str(old.get("email") or "").lower()):
            history = old.get("submission_history", [])
            history.append({
                "timestamp": old.get("last_updated", "未知"),
                "resume_file_path": old.get("resume_file_path", ""),
                "portfolio_file_path": old.get("portfolio_file_path", ""),
            })
            new_candidate["submission_history"] = history
            old_candidates[i] = new_candidate
            return old_candidates, True
    new_candidate.setdefault("submission_history", [])
    old_candidates.append(new_candidate)
    return old_candidates, False


# ══════════════════════════════════════════════════════════════════════════════
# 八、Excel 导出（新增获奖/GPA/个人特质/评分摘要列）
# ══════════════════════════════════════════════════════════════════════════════

_HDR_FILL = PatternFill("solid", fgColor="1F3864")
_HDR_FONT = Font(name="Arial", color="FFFFFF", bold=True, size=11)
_DUP_FILL = PatternFill("solid", fgColor="C00000")
_DUP_FONT = Font(name="Arial", color="FFFFFF", bold=True, size=11)
_BODY_FONT = Font(name="Arial", size=10)
_BOLD_FONT = Font(name="Arial", bold=True, size=10)
_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
_LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)
_THIN = Side(style="thin", color="CCCCCC")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)


def _write_header(ws, headers: list, fill, font) -> None:
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = fill
        cell.font = font
        cell.alignment = _CENTER
        cell.border = _BORDER
    ws.row_dimensions[1].height = 28


def _auto_width(ws, headers: list, max_width: int = 45) -> None:
    for col_idx in range(1, len(headers) + 1):
        col_letter = get_column_letter(col_idx)
        max_len = max(
            (len(str(c.value or "")) for c in ws[col_letter]),
            default=len(headers[col_idx - 1])
        )
        ws.column_dimensions[col_letter].width = min(max_len * 1.2 + 2, max_width)


def _file_uri(fpath: str) -> str:
    if not fpath:
        return ""
    p = Path(fpath)
    if not p.exists():
        return ""
    try:
        return p.resolve().as_uri()
    except Exception:
        return ""


def _write_file_cell(cell, fpath: str, bold: bool = False) -> None:
    if not fpath:
        cell.value = ""
        return
    uri = _file_uri(fpath)
    cell.value = Path(fpath).name
    if uri:
        cell.hyperlink = uri
    cell.font = Font(
        name="Arial", size=10, bold=bold,
        color="0563C1" if uri else "000000",
        underline="single" if uri else None,
    )


def save_to_excel(candidates: list, path) -> None:
    """
    导出 Excel：
    - Sheet「全部候选人」：按岗位分组 + 组内评分降序，岗位间有分隔行
    - Sheet「各岗位排名」（每个岗位一个 Sheet）：只含该岗位候选人，评分降序
    - Sheet「重复投递记录」：多次投递者的历史
    """
    wb = openpyxl.Workbook()

    # 岗位分隔行样式
    _SEP_FILL = PatternFill("solid", fgColor="D9E1F2")
    _SEP_FONT = Font(name="Arial", bold=True, size=11, color="1F3864")

    main_headers = [
        "排名", "姓名", "性别", "电话", "邮箱", "毕业院校", "专业", "学历", "意向岗位",
        "编程技能", "办公软件", "AI工具", "英语水平",
        "实习经历", "校园领导经历", "获奖情况", "GPA/成绩",
        "个人优势（表单）", "个人特质（简历提取）",
        "AI综合评分(/100)", "评分摘要",
        "简历文件", "作品集",
    ]

    def _write_candidate_row(ws, row_i: int, c: dict, rank: str = "") -> None:
        score_data   = c.get("jd_score", {})
        score_total  = score_data.get("total", "")
        score_summary = score_data.get("summary", "")
        values = [
            rank,
            c.get("name", ""),
            c.get("gender", ""),
            c.get("phone", ""),
            c.get("email", ""),
            c.get("school", ""),
            c.get("major", ""),
            c.get("degree", ""),
            c.get("target_position", ""),
            ", ".join(c.get("skills_programming") or []),
            ", ".join(c.get("skills_office")      or []),
            ", ".join(c.get("skills_ai_tools")    or []),
            c.get("english_level") or "",
            c.get("internship")    or "",
            c.get("campus_leadership") or "",
            c.get("awards") or "",
            c.get("gpa")    or "",
            (", ".join(c.get("strengths") or [])
             if isinstance(c.get("strengths"), list)
             else str(c.get("strengths") or "")),
            ", ".join(c.get("personal_traits") or []),
            score_total,
            score_summary,
        ]
        center_cols = {1, 3, 8, 9, 20}
        for col_i, val in enumerate(values, 1):
            cell = ws.cell(row=row_i, column=col_i, value=str(val) if val is not None else "")
            cell.border    = _BORDER
            cell.alignment = _CENTER if col_i in center_cols else _LEFT
            cell.font      = _BODY_FONT
        # 文件列（第22、23列）
        for col_i, fpath_key in [(22, "resume_file_path"), (23, "portfolio_file_path")]:
            cell = ws.cell(row=row_i, column=col_i)
            cell.border = _BORDER; cell.alignment = _LEFT
            _write_file_cell(cell, c.get(fpath_key, ""))

    # ── Sheet 1：全部候选人（按岗位分组 + 组内评分降序）────────────────────
    ws1 = wb.active
    ws1.title        = "全部候选人"
    ws1.freeze_panes = "B2"
    _write_header(ws1, main_headers, _HDR_FILL, _HDR_FONT)

    # 按岗位分组，每组内按评分降序
    from collections import defaultdict
    groups: dict[str, list] = defaultdict(list)
    for c in candidates:
        pos = c.get("target_position") or "其他"
        groups[pos].append(c)
    for pos in groups:
        groups[pos].sort(
            key=lambda x: x.get("jd_score", {}).get("total", -1) or -1,
            reverse=True,
        )

    row_i = 2
    for pos, group in sorted(groups.items()):
        # 岗位分隔行
        sep_cell = ws1.cell(row=row_i, column=1,
                            value=f"▌ {pos}  （共 {len(group)} 人）")
        sep_cell.fill = _SEP_FILL
        sep_cell.font = _SEP_FONT
        sep_cell.alignment = _LEFT
        ws1.merge_cells(start_row=row_i, start_column=1,
                        end_row=row_i, end_column=len(main_headers))
        ws1.row_dimensions[row_i].height = 22
        row_i += 1

        for rank_i, c in enumerate(group, 1):
            _write_candidate_row(ws1, row_i, c, rank=str(rank_i))
            row_i += 1

    ws1.auto_filter.ref = f"A1:{get_column_letter(len(main_headers))}1"
    _auto_width(ws1, main_headers)

    # ── 各岗位独立 Sheet（评分降序）──────────────────────────────────────────
    for pos, group in sorted(groups.items()):
        safe_name = re.sub(r'[\\/*?\[\]:]', '_', pos)[:28]  # Excel sheet名限31字符
        ws = wb.create_sheet(safe_name)
        ws.freeze_panes = "B2"
        _write_header(ws, main_headers, _HDR_FILL, _HDR_FONT)
        for rank_i, c in enumerate(group, 1):
            _write_candidate_row(ws, rank_i + 1, c, rank=str(rank_i))
        ws.auto_filter.ref = f"A1:{get_column_letter(len(main_headers))}1"
        _auto_width(ws, main_headers)

    # ── Sheet「重复投递记录」────────────────────────────────────────────────
    ws2 = wb.create_sheet("重复投递记录")
    ws2.freeze_panes = "A2"

    dup_headers = ["姓名", "电话", "版本", "投递时间", "简历文件", "作品集"]
    _write_header(ws2, dup_headers, _DUP_FILL, _DUP_FONT)

    dup_row = 2
    for c in candidates:
        history = c.get("submission_history", [])
        if not history:
            continue

        def _fill_dup_row(ri, name, phone, ver, ts, rp, pp, bold=False):
            for ci, v in enumerate([name, phone, ver, ts], 1):
                cell = ws2.cell(row=ri, column=ci, value=v or "")
                cell.border = _BORDER; cell.alignment = _CENTER
                cell.font   = _BOLD_FONT if bold else _BODY_FONT
            for ci, fpath in [(5, rp), (6, pp)]:
                cell = ws2.cell(row=ri, column=ci)
                cell.border = _BORDER; cell.alignment = _LEFT
                _write_file_cell(cell, fpath, bold=bold)

        _fill_dup_row(dup_row,
                      c.get("name",""), c.get("phone",""), "最新", c.get("last_updated",""),
                      c.get("resume_file_path",""), c.get("portfolio_file_path",""), bold=True)
        dup_row += 1
        for h in history:
            _fill_dup_row(dup_row,
                          c.get("name",""), c.get("phone",""), "历史", h.get("timestamp",""),
                          h.get("resume_file_path",""), h.get("portfolio_file_path",""))
            dup_row += 1

    ws2.auto_filter.ref = f"A1:{get_column_letter(len(dup_headers))}1"
    _auto_width(ws2, dup_headers)

    wb.save(str(path))
